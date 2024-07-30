import os
from docx import Document
from docx.oxml.ns import qn

def convert_docx_to_markdown(docx_path, output_folder):
    # Create output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Load the .docx document
    doc = Document(docx_path)
    
    markdown_lines = []
    image_counter = 1

    for element in doc.element.body:
        if element.tag == qn('w:p'):
            # Paragraph
            paragraph = element
            para_obj = doc.paragraphs[doc.element.body.index(paragraph)]
            if para_obj.style.name.startswith('Heading'):
                level = int(para_obj.style.name[-1])
                markdown_lines.append('#' * level + ' ' + para_obj.text)
            else:
                markdown_lines.append(para_obj.text)
        elif element.tag == qn('w:tbl'):
            # Table
            table = doc.tables[doc.element.body.index(element)]
            markdown_lines.append('\n')
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                markdown_lines.append('| ' + ' | '.join(row_data) + ' |')
                markdown_lines.append('|' + '---|' * len(row_data))
            markdown_lines.append('\n')
        elif element.tag == qn('w:drawing') or element.tag == qn('w:pict'):
            # Image
            for rel in doc.part.rels:
                if "image" in doc.part.rels[rel].target_ref:
                    img = doc.part.rels[rel].target_part.blob
                    img_filename = f'image_{image_counter}.png'
                    img_path = os.path.join(output_folder, img_filename)
                    with open(img_path, 'wb') as img_file:
                        img_file.write(img)
                    markdown_lines.append(f'![Image {image_counter}]({img_filename})')
                    image_counter += 1

    # Write the markdown content to a file
    markdown_path = os.path.join(output_folder, 'output.md')
    with open(markdown_path, 'w') as md_file:
        md_file.write('\n'.join(markdown_lines))

# Example usage
docx_path = "D:/OneDrive - IZ/MxIoT-Documents/Projects/GenAI-Intelizign/RequirementGenerator/ExampleDocs/appendix_6m_tool_requirement_specifications_v2_project_portfolio_management_tool.docx"  # Replace with your input docx file path
output_folder = "D:/OneDrive - IZ/MxIoT-Documents/Projects/GenAI-Intelizign/RequirementGenerator/ExampleDocs"
convert_docx_to_markdown(docx_path, output_folder)
